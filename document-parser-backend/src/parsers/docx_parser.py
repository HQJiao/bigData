from pathlib import Path
from typing import List
from docx import Document as DocxDocument

from src.parsers.base import IParser
from src.parsers.models import ParsedContent
from src.parsers.registry import parser_registry


class DocxParser(IParser):
    def parse(self, file_path: Path) -> ParsedContent:
        doc = DocxDocument(file_path)
        
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables_text.append(row_text)
        
        all_text = "\n".join(paragraphs + tables_text)
        
        return ParsedContent(
            text=all_text,
            parser_type="docx",
            metadata={"paragraph_count": len(paragraphs), "table_count": len(doc.tables)},
            table_count=len(doc.tables),
        )

    def supported_extensions(self) -> List[str]:
        return [".docx"]


parser_registry.register(DocxParser())
