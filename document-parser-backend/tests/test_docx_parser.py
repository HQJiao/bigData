import pytest
import tempfile
from pathlib import Path

from src.parsers.docx_parser import DocxParser


def test_docx_parser_supported_extensions():
    parser = DocxParser()
    assert ".docx" in parser.supported_extensions()


def test_docx_parser_basic():
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    parser = DocxParser()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Test Content")
        doc.save(f.name)
        temp_path = Path(f.name)

    try:
        result = parser.parse(temp_path)
        assert result.parser_type == "docx"
        assert "Hello World" in result.text
    finally:
        temp_path.unlink()


def test_docx_parser_metadata():
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    parser = DocxParser()
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(f.name)
        temp_path = Path(f.name)

    try:
        result = parser.parse(temp_path)
        assert "paragraph_count" in result.metadata
        assert "table_count" in result.metadata
    finally:
        temp_path.unlink()
